import docx
import json
from datetime import datetime
from dateutil import parser
from pdfkit import pdfkit
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from icecream import ic


def is_date(value):
    try:
        parser.parse(value)
        return True
    except ValueError:
        return False


def replace_text_in_docx(docx_path, data):
    doc = docx.Document(docx_path)
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            # ic(value, key)
            if isinstance(value, list):
                break
                # Handle array block
                # ic("array")
                start_tag = f"{{{key.upper()}}}"
                end_tag = f"{{/{key.upper()}}}"
                if start_tag in paragraph.text:
                    for paragraph_index, paragraph in enumerate(doc.paragraphs):
                        # rest of the code remains the same doc.paragraphs.index(paragraph)
                        for item in value:
                            new_paragraph = doc.paragraphs[
                                paragraph_index
                            ].insert_paragraph_after(paragraph.text)
                            for run in new_paragraph.runs:
                                new_paragraph.text = new_paragraph.text.replace(
                                    "[STUL]", str(item["stul"])
                                )
                            new_paragraph.text = new_paragraph.text.replace(
                                start_tag, ""
                            ).replace(end_tag, "")
                        doc.paragraphs[paragraph_index].clear()
            else:
                # Handle key-value replacement
                # ic("key-value")
                placeholder = f"[{key.upper()}]"
                if placeholder in paragraph.text:
                    # Format the datum_narozeni field to "DD.MM.YYYY"
                    if is_date(value):
                        value = datetime.strptime(value, "%Y-%m-%d").strftime(
                            "%d.%m.%Y"
                        )

                    # Replace the placeholder with the value
                    replaced_text = str(value)
                    paragraph.text = paragraph.text.replace(placeholder, replaced_text)

                    # Track the start and end positions of the replaced text
                    start_index = paragraph.text.find(replaced_text)
                    end_index = start_index + len(replaced_text)

                    # Iterate over runs to apply bold formatting within the replaced text range
                    for run in paragraph.runs:
                        run_text = run.text
                        if replaced_text in run_text:
                            ic(str(value), run.text)
                            run_start_index = run_text.find(replaced_text)
                            # run_text_start = paragraph.text.find(run.text)
                            run_end_index = run_start_index + len(replaced_text)
                            ic(start_index, end_index)
                            ic(
                                run_start_index,
                                run_end_index,
                            )
                            # Check if the run's text is within the replaced text range
                            if (
                                run_start_index <= start_index
                                and run_end_index >= end_index
                            ):
                                # Apply bold formatting to the specified portion of the run's text
                                run = run.clear()
                                run.add_text(run_text[:run_start_index])
                                run.add_text(
                                    f"{run_text[run_start_index:end_index]}", bold=True
                                )
                                run.add_text(run_text[end_index:])

    return doc


def main():
    # Path to the Word document
    docx_path = r"D:\\soubor1.docx"

    # Path to the JSON file
    json_path = r"d:\\data.json"

    # Load data from the JSON file
    with open(json_path, "r") as json_file:
        data = json.load(json_file)

    # Replace placeholders in the Word document
    doc = replace_text_in_docx(docx_path, data)

    # Generate timestamp
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

    # New file name with timestamp
    new_docx_path = rf"D:/novy_soubor_{timestamp}.docx"
    new_pdf_path = rf"D:/novy_soubor_{timestamp}.pdf"

    # Save the new document in DOCX format
    doc.save(new_docx_path)

    # Convert the DOCX document to PDF
    # convert_to_pdf(new_docx_path, new_pdf_path)


def convert_to_pdf(docx_path, pdf_path):
    pdfkit.from_file(docx_path, pdf_path)


if __name__ == "__main__":
    main()
