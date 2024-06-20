import pdfplumber
from docx import Document
from docx.shared import Inches
import re


def is_header_or_footer(text, page_texts):
    specific_strings = [
        "BABOK® Guide, Version 2.0",
        "Order ID: IIBA-200911231134-455082",
        "Licensed to Gustavo Simues <gustavo.simoes@fattocs.com.br>",
    ]
    specific_contains_strings = ["A Guide to the Business Analysis Body of Knowledge"]
    if any(text.startswith(s) for s in specific_strings):
        return True
    if any(s in text for s in specific_contains_strings):
        return True
    # Simple heuristic to determine if text is header or footer:
    # If it appears on every page in the same position.
    return text in page_texts and len(page_texts[text]) == len(page_texts)


def extract_text_from_pdf(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        page_texts = {}
        all_text = []

        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if not page_text:
                continue

            lines = page_text.split("\n")
            for line in lines:
                if line not in page_texts:
                    page_texts[line] = []
                page_texts[line].append(i)
            all_text.append(lines)

        # Determine headers and footers
        headers_and_footers = [
            text for text in page_texts if is_header_or_footer(text, page_texts)
        ]
        clean_text = []

        for lines in all_text:
            page_clean_text = [
                line for line in lines if line not in headers_and_footers
            ]
            clean_text.extend(page_clean_text)

        return "\n".join(clean_text)


def convert_pdf_to_docx(pdf_path, docx_path):
    text = extract_text_from_pdf(pdf_path)

    document = Document()

    for line in text.split("\n"):
        document.add_paragraph(line)
    document.save(docx_path)
    exit(0)
    # For adding images, you need to extract and save them first, then add to docx
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages):
            for img in page.images:
                # Extract image
                x0, y0, x1, y1 = img["x0"], img["top"], img["x1"], img["bottom"]
                image = page.within_bbox((x0, y0, x1, y1)).to_image()

                # Save image
                image_path = f"image_{i}_{x0}_{y0}.png"
                image.save(image_path, format="PNG")

                # Add image to the document
                document.add_picture(image_path, width=Inches(5))

    document.save(docx_path)


# Usage example
# PDF_PATH = "e:/Můj disk/data-modelovani/analýza-architektura/Babok Guide.pdf"
PDF_PATH = "D:/Users/PWeigner/Downloads/form4data.pdf"
# DOCX_PATH = "d:/Babok Guide.docx"
DOCX_PATH = "d:/form4data.docx"
convert_pdf_to_docx(pdf_path=PDF_PATH, docx_path=DOCX_PATH)
